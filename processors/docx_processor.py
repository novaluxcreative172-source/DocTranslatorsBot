"""
processors/docx_processor.py
Translates a .docx file while preserving structure: paragraphs, runs
(so bold/italic/font stays intact), tables, and headers/footers.

Strategy: translate each run's text in place rather than the whole
paragraph at once. This keeps inline formatting (e.g. "this is **bold**
text") attached to the correct words instead of collapsing into one
plain-text blob.
"""

from docx import Document
from translator import translate_text


def _translate_paragraph(paragraph, target_lang):
    for run in paragraph.runs:
        if run.text and run.text.strip():
            run.text = translate_text(run.text, target_lang)


def _translate_table(table, target_lang):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _translate_paragraph(paragraph, target_lang)


def process(input_path: str, output_path: str, target_lang: str) -> None:
    doc = Document(input_path)

    for paragraph in doc.paragraphs:
        _translate_paragraph(paragraph, target_lang)

    for table in doc.tables:
        _translate_table(table, target_lang)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _translate_paragraph(paragraph, target_lang)
        for paragraph in section.footer.paragraphs:
            _translate_paragraph(paragraph, target_lang)

    doc.save(output_path)
